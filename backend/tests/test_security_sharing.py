import pytest
import io
import uuid
import docx
import openpyxl
from pypdf import PdfWriter
from sqlalchemy import select, text
from app.models import User, Folder, FileItem, Workspace, WorkspaceMember
from app.services.chunking_service import chunking_service
from app.services.document_service import document_service
from app.services.access_service import access_service
from app.services.search_service import search_service
from app.schemas.search import SearchRequest

@pytest.mark.asyncio
async def test_multi_format_extraction():
    """Test Word, Excel, and PDF text extraction."""
    # 1. Word DOCX
    doc = docx.Document()
    doc.add_heading("2026 AI 연구 전략 보고서", level=1)
    doc.add_paragraph("본 문서는 멀티모달 에이전트와 지식 저장소 통합에 대한 핵심 전략입니다.")
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_bytes = docx_buf.getvalue()

    docx_text = chunking_service.extract_text_from_docx(docx_bytes)
    assert "2026 AI 연구 전략 보고서" in docx_text
    assert "멀티모달 에이전트" in docx_text

    # 2. Excel XLSX
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "예산안"
    ws.append(["항목", "금액", "담당자"])
    ws.append(["GPU 클러스터 대여", "50000000", "우광완"])
    ws.append(["스토리지 인프라", "12000000", "김연구"])
    xlsx_buf = io.BytesIO()
    wb.save(xlsx_buf)
    xlsx_bytes = xlsx_buf.getvalue()

    xlsx_text = chunking_service.extract_text_from_excel(xlsx_bytes)
    assert "[Sheet: 예산안]" in xlsx_text
    assert "GPU 클러스터 대여" in xlsx_text
    assert "우광완" in xlsx_text

    # 3. PDF
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    pdf_buf = io.BytesIO()
    writer.write(pdf_buf)
    pdf_bytes = pdf_buf.getvalue()
    pdf_text = chunking_service.extract_text_from_pdf(pdf_bytes)
    assert isinstance(pdf_text, str)

@pytest.mark.asyncio
async def test_workspace_and_permission_isolation(db_session):
    """
    Test Workspace Architecture:
    1. Alice creates Workspace A (Engineering) and Workspace B (Design) - owning multiple workspaces.
    2. Bob is invited to Workspace A.
    3. Charlie is invited to Workspace B only.
    4. Alice creates confidential files in Workspace A and B.
    5. Verify Bob can access Workspace A files, but CANNOT access Workspace B files.
    6. Verify Charlie can access Workspace B files, but CANNOT access Workspace A files.
    7. Remove Bob from Workspace A -> Verify Bob's access is immediately revoked!
    """
    uid = str(uuid.uuid4())[:8]
    alice = User(email=f"alice_{uid}@project.run", name="Alice (오너)", is_admin=False, is_approved=True)
    bob = User(email=f"bob_{uid}@project.run", name="Bob (엔지니어)", is_admin=False, is_approved=True)
    charlie = User(email=f"charlie_{uid}@project.run", name="Charlie (디자이너)", is_admin=False, is_approved=True)
    db_session.add_all([alice, bob, charlie])
    await db_session.commit()
    for u in [alice, bob, charlie]:
        await db_session.refresh(u)

    # 1. Alice creates Workspace A (Engineering) and Workspace B (Design)
    ws_a = Workspace(
        name="엔지니어링 팀",
        slug=f"engineering-{uid}",
        owner_id=alice.id,
        icon="code"
    )
    ws_b = Workspace(
        name="디자인 팀",
        slug=f"design-{uid}",
        owner_id=alice.id,
        icon="palette"
    )
    db_session.add_all([ws_a, ws_b])
    await db_session.commit()
    for w in [ws_a, ws_b]:
        await db_session.refresh(w)

    # Add Alice as owner of both workspaces
    mem_alice_a = WorkspaceMember(workspace_id=ws_a.id, user_id=alice.id, role="owner")
    mem_alice_b = WorkspaceMember(workspace_id=ws_b.id, user_id=alice.id, role="owner")
    # Add Bob to Workspace A
    mem_bob_a = WorkspaceMember(workspace_id=ws_a.id, user_id=bob.id, role="member")
    # Add Charlie to Workspace B
    mem_charlie_b = WorkspaceMember(workspace_id=ws_b.id, user_id=charlie.id, role="member")
    
    db_session.add_all([mem_alice_a, mem_alice_b, mem_bob_a, mem_charlie_b])
    await db_session.commit()

    # 2. Files in Workspace A and Workspace B
    file_a = FileItem(
        workspace_id=ws_a.id,
        created_by=alice.id,
        name="Backend_Architecture_2026.md",
        file_type="note",
        content="FastAPI와 pgvector를 활용한 고성능 백엔드 설계 명세서입니다.",
        size_bytes=100,
        is_markdown=True
    )
    file_b = FileItem(
        workspace_id=ws_b.id,
        created_by=alice.id,
        name="Design_System_Figma.md",
        file_type="note",
        content="Figma 기반 다크 테마 및 유리모피즘 디자인 가이드입니다.",
        size_bytes=100,
        is_markdown=True
    )
    db_session.add_all([file_a, file_b])
    await db_session.commit()
    for f in [file_a, file_b]:
        await db_session.refresh(f)

    # Index document chunks for search
    await document_service.index_file_chunks(db_session, file_a)
    await document_service.index_file_chunks(db_session, file_b)

    # 3. Verify Workspace Isolation
    # Alice (owner of both) can access both
    assert await access_service.can_access_file(db_session, alice, file_a.id) is True
    assert await access_service.can_access_file(db_session, alice, file_b.id) is True

    # Bob (member of A) can access file_a, but NOT file_b
    assert await access_service.can_access_file(db_session, bob, file_a.id) is True
    assert await access_service.can_access_file(db_session, bob, file_b.id) is False

    # Charlie (member of B) can access file_b, but NOT file_a
    assert await access_service.can_access_file(db_session, charlie, file_a.id) is False
    assert await access_service.can_access_file(db_session, charlie, file_b.id) is True

    # 4. Search Isolation: Search file_a content
    # Bob searching in workspace_a finds file_a
    search_req_a = SearchRequest(query="백엔드 설계", workspace_id=ws_a.id, include_keyword_match=True)
    bob_res_a = await search_service.search(db_session, search_req_a, current_user=bob)
    assert bob_res_a.total_results >= 1
    assert any(r.file_id == file_a.id for r in bob_res_a.results)

    # Charlie searching in workspace_a gets 0 results (unauthorized workspace)
    charlie_res_a = await search_service.search(db_session, search_req_a, current_user=charlie)
    assert charlie_res_a.total_results == 0

    # 5. Remove Bob from Workspace A -> Verify immediate revocation
    await db_session.delete(mem_bob_a)
    await db_session.commit()

    assert await access_service.can_access_file(db_session, bob, file_a.id) is False
    bob_res_revoked = await search_service.search(db_session, search_req_a, current_user=bob)
    assert bob_res_revoked.total_results == 0

    # Cleanup
    await db_session.execute(text(f"DELETE FROM kb_users WHERE id IN ('{alice.id}', '{bob.id}', '{charlie.id}')"))
    await db_session.commit()

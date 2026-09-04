import re
import io
from typing import List, Dict, Optional
from pypdf import PdfReader
import docx
import openpyxl

class ChunkingService:
    @staticmethod
    def chunk_markdown(content: str, max_chunk_size: int = 800, overlap: int = 100) -> List[Dict[str, any]]:
        """
        Split markdown content into semantic chunks based on headers (#, ##, ###) and paragraphs.
        Keeps header context attached to child chunks for higher retrieval quality.
        """
        if not content or not content.strip():
            return []

        lines = content.splitlines()
        chunks = []
        
        current_headers = {1: "", 2: "", 3: "", 4: ""}
        current_chunk_lines = []
        current_chunk_len = 0

        def get_header_context():
            hdrs = [v for k, v in sorted(current_headers.items()) if v]
            return " > ".join(hdrs)

        def flush_chunk():
            nonlocal current_chunk_lines, current_chunk_len
            if not current_chunk_lines:
                return
            
            chunk_body = "\n".join(current_chunk_lines).strip()
            if not chunk_body:
                current_chunk_lines = []
                current_chunk_len = 0
                return

            header_ctx = get_header_context()
            if header_ctx and not chunk_body.startswith("#"):
                full_text = f"[{header_ctx}]\n{chunk_body}"
            else:
                full_text = chunk_body

            chunks.append({
                "chunk_index": len(chunks),
                "content": full_text
            })
            
            # Carry over overlap if needed
            if overlap > 0 and len(current_chunk_lines) > 2:
                current_chunk_lines = current_chunk_lines[-2:]
                current_chunk_len = sum(len(l) for l in current_chunk_lines)
            else:
                current_chunk_lines = []
                current_chunk_len = 0

        for line in lines:
            header_match = re.match(r"^(#{1,4})\s+(.+)$", line.strip())
            if header_match:
                # If we encounter a new header and have accumulated content, flush it
                if current_chunk_lines:
                    flush_chunk()
                
                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                
                current_headers[level] = title
                # Clear lower level headers
                for l in range(level + 1, 5):
                    current_headers[l] = ""
                
                current_chunk_lines.append(line)
                current_chunk_len += len(line)
            else:
                current_chunk_lines.append(line)
                current_chunk_len += len(line)
                
                if current_chunk_len >= max_chunk_size:
                    flush_chunk()

        if current_chunk_lines:
            flush_chunk()

        return chunks

    @staticmethod
    def chunk_plain_text(content: str, max_chunk_size: int = 800, overlap: int = 100) -> List[Dict[str, any]]:
        """Split plain text into paragraph-aware chunks."""
        if not content or not content.strip():
            return []

        paragraphs = content.split("\n\n")
        chunks = []
        current_chunk = []
        current_len = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if current_len + len(para) > max_chunk_size and current_chunk:
                chunks.append({
                    "chunk_index": len(chunks),
                    "content": "\n\n".join(current_chunk)
                })
                current_chunk = []
                current_len = 0

            current_chunk.append(para)
            current_len += len(para)

        if current_chunk:
            chunks.append({
                "chunk_index": len(chunks),
                "content": "\n\n".join(current_chunk)
            })

        return chunks

    @staticmethod
    def sanitize_text(text: str) -> str:
        """
        Text as a database column can actually hold it.

        A PDF's extracted text can carry a NUL byte — this one had two, in the
        middle of a UUID whose hyphens had been mangled — and Postgres refuses
        any text containing 0x00 outright ("invalid byte sequence for encoding
        UTF8"), which failed the insert, broke the transaction, and took the
        whole upload down with it. Lone surrogates come out of the same kind of
        salvaged text and are refused the same way, so the round trip through
        UTF-8 drops those too.
        """
        if not text:
            return text
        without_nul = text.replace("\x00", "")
        return without_nul.encode("utf-8", "ignore").decode("utf-8", "ignore")

    @staticmethod
    def extract_text_from_pdf(pdf_bytes: bytes) -> str:
        """Extract plain text from PDF binary data."""
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            text_pages = []
            for i, page in enumerate(reader.pages):
                t = page.extract_text()
                if t and t.strip():
                    text_pages.append(f"### [PDF Page {i+1}]\n{t.strip()}")
            return ChunkingService.sanitize_text("\n\n".join(text_pages))
        except Exception as e:
            print(f"[PDF Extract Warning] Failed to extract text from PDF: {e}")
            return ""

    @staticmethod
    def extract_text_from_docx(docx_bytes: bytes) -> str:
        """Extract text, headers, and tables from Word (.docx) documents."""
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))
            lines = []

            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                if p.style and p.style.name.startswith("Heading 1"):
                    lines.append(f"# {text}")
                elif p.style and p.style.name.startswith("Heading 2"):
                    lines.append(f"## {text}")
                elif p.style and p.style.name.startswith("Heading 3"):
                    lines.append(f"### {text}")
                else:
                    lines.append(text)

            # Extract tables
            for t_idx, table in enumerate(doc.tables):
                lines.append(f"\n### [Table {t_idx + 1}]")
                for row_idx, row in enumerate(table.rows):
                    row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    row_str = " | ".join(row_cells)
                    lines.append(f"| {row_str} |")
                    if row_idx == 0:
                        separator = " | ".join(["---"] * len(row_cells))
                        lines.append(f"| {separator} |")

            return ChunkingService.sanitize_text("\n\n".join(lines))
        except Exception as e:
            print(f"[Word Extract Warning] Failed to extract text from DOCX: {e}")
            return ""

    @staticmethod
    def extract_text_from_excel(excel_bytes: bytes) -> str:
        """Extract sheets, column headers, and row data from Excel (.xlsx, .xls) files."""
        try:
            wb = openpyxl.load_workbook(io.BytesIO(excel_bytes), data_only=True)
            sheets_text = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_lines = [f"## [Sheet: {sheet_name}]"]

                rows = list(sheet.iter_rows(values_only=True))
                if not rows:
                    continue

                for r_idx, row in enumerate(rows[:500]): # Limit to first 500 rows per sheet
                    # Filter out all-None rows
                    if not any(cell is not None for cell in row):
                        continue
                    
                    row_cells = [str(cell).strip() if cell is not None else "" for cell in row]
                    # Trim empty trailing cells
                    while row_cells and not row_cells[-1]:
                        row_cells.pop()

                    if row_cells:
                        sheet_lines.append(" | ".join(row_cells))

                sheets_text.append("\n".join(sheet_lines))

            return ChunkingService.sanitize_text("\n\n".join(sheets_text))
        except Exception as e:
            print(f"[Excel Extract Warning] Failed to extract text from Excel: {e}")
            return ""

chunking_service = ChunkingService()
